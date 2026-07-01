# report_generator.py
"""
Generates the quarterly market report deliverable for the MOG division
(Individual KPI: "Publish Market reports quarterly for MOG division").

The report pulls directly from the SQLite store via analytics.py so every
figure ties back to logged benchmark/product prices and monitored market
developments — no numbers are invented at report time.
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import REPORTS_DIR, MOG_DIVISION_NAME
from database import log_quarterly_report
from analytics import (
    quarterly_benchmark_stats,
    quarterly_product_stats,
    developments_for_period,
)

GOLD = RGBColor(0xB8, 0x8A, 0x1E)
DARK = RGBColor(0x1A, 0x1F, 0x2B)
GREY = RGBColor(0x5A, 0x66, 0x78)


def _set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, color=GOLD)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            _set_cell_text(cells[i], "" if val is None else val)
    return table


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK
    return h


def generate_quarterly_report(crude_df, product_df, dev_df, year: int, quarter: str,
                               outlook_notes: str = "", generated_by: str = "MOG Analyst") -> str:
    """Builds the quarterly Word report and returns the saved file path."""
    b_stats = quarterly_benchmark_stats(crude_df, year, quarter)
    p_stats = quarterly_product_stats(product_df, year, quarter)
    devs = developments_for_period(dev_df, year, quarter)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{quarter} {year} Market Intelligence Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(MOG_DIVISION_NAME)
    run.font.size = Pt(13)
    run.font.color.rgb = GOLD

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated {datetime.now().strftime('%d %B %Y')} • Prepared by {generated_by}")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY

    doc.add_paragraph()

    # Executive summary
    _heading(doc, "Executive Summary", level=1)
    if not b_stats.empty:
        movers = b_stats.sort_values("Change %", ascending=False, na_position="last")
        best = movers.iloc[0]
        worst = movers.iloc[-1]
        summary_text = (
            f"During {quarter} {year}, {best['Benchmark']} was the strongest-performing tracked "
            f"benchmark ({'+' if best['Change %'] and best['Change %'] >= 0 else ''}{best['Change %']}% quarter-on-quarter), "
            f"while {worst['Benchmark']} moved {worst['Change %']}%. "
            f"{len(devs)} market development(s) were logged in the monitoring feed this quarter."
        )
    else:
        summary_text = (
            f"No benchmark price readings were recorded for {quarter} {year} at the time of "
            f"generation. Populate the pipeline or enter manual readings, then regenerate this report."
        )
    doc.add_paragraph(summary_text)

    # Section I: Crude benchmark review
    _heading(doc, "I. Crude Benchmark Price Review", level=1)
    doc.add_paragraph(
        "Quarter-over-quarter movement in Brent, WTI, Dubai, Oman, and Kuwait Export Crude, "
        "based on daily readings logged in the market intelligence system."
    )
    if not b_stats.empty:
        _add_table(
            doc,
            ["Benchmark", "Open", "Close", "High", "Low", "Average", "Change", "Change %", "Readings"],
            b_stats.values.tolist(),
        )
    else:
        doc.add_paragraph("No data available for this period.")

    doc.add_paragraph()

    # Section II: Refined products
    _heading(doc, "II. Refined Product Proxy Review", level=1)
    doc.add_paragraph(
        "Singapore / regional refined product proxy movement (Naphtha, Gasoline, Jet/Kerosene, "
        "Gasoil, Fuel Oil, LPG)."
    )
    if not p_stats.empty:
        _add_table(
            doc,
            ["Product", "Open", "Close", "High", "Low", "Average", "Change", "Change %", "Readings"],
            p_stats.values.tolist(),
        )
    else:
        doc.add_paragraph("No data available for this period.")

    doc.add_paragraph()

    # Section III: Market developments
    _heading(doc, "III. Major Market Developments", level=1)
    doc.add_paragraph(
        "Developments monitored across OPEC+, IEA, EIA publications, and geopolitical, "
        "economic, and industry news impacting crude and refined product markets this quarter."
    )
    if not devs.empty:
        for category, group in devs.groupby("category"):
            _heading(doc, category, level=2)
            for _, row in group.iterrows():
                p = doc.add_paragraph(style="List Bullet")
                date_str = row.get("dev_date", "")
                impact = row.get("impact", "")
                run = p.add_run(f"[{date_str}] ")
                run.font.color.rgb = GREY
                run.font.size = Pt(9)
                p.add_run(f"{row.get('headline', '')} ")
                impact_run = p.add_run(f"(Impact: {impact})")
                impact_run.italic = True
                impact_run.font.size = Pt(9)
                impact_run.font.color.rgb = GOLD
    else:
        doc.add_paragraph("No developments were logged for this period.")

    doc.add_paragraph()

    # Section IV: Outlook / analyst notes
    _heading(doc, "IV. Outlook & Analyst Notes", level=1)
    doc.add_paragraph(outlook_notes.strip() if outlook_notes.strip() else
                       "No additional analyst commentary was provided for this quarter.")

    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("KNPC Market Intelligence & Business Data Platform — Confidential, Internal Use Only")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"MOG_Quarterly_Market_Report_{quarter}_{year}.docx"
    output_path = REPORTS_DIR / filename
    doc.save(output_path)

    log_quarterly_report(quarter=quarter, year=year, file_path=str(output_path), generated_by=generated_by)
    return str(output_path)


def list_generated_reports():
    """Returns a sorted list of previously generated report file paths."""
    if not REPORTS_DIR.exists():
        return []
    return sorted(REPORTS_DIR.glob("*.docx"), key=lambda f: f.stat().st_mtime, reverse=True)
